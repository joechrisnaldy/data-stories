"""Convert the essay draft (simple markdown) to a Word .docx for review.

Handles the constructs this draft actually uses: #/## headings, paragraphs,
![img](path), > blockquote, - bullets, --- rule, inline **bold** / *italic*
and [text](url) rendered as text.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor

INLINE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|\[.+?\]\(.+?\))")
LINK = re.compile(r"\[(.+?)\]\((.+?)\)")


def add_runs(par, text):
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            par.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*"):
            par.add_run(part[1:-1]).italic = True
        elif LINK.fullmatch(part):
            par.add_run(LINK.fullmatch(part).group(1))
        else:
            par.add_run(part)


def convert(md_path, docx_path):
    md_path = Path(md_path)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for block in md_path.read_text().split("\n\n"):
        block = block.strip()
        if not block or block == "---":
            continue
        if block.startswith("# "):
            doc.add_heading(block[2:], level=0)
        elif block.startswith("## "):
            doc.add_heading(block[3:], level=1)
        elif block.startswith("!["):
            m = re.match(r"!\[.*?\]\((.+?)\)", block)
            img = (md_path.parent / m.group(1)).resolve()
            doc.add_picture(str(img), width=Inches(6.2))
        elif block.startswith("> "):
            par = doc.add_paragraph()
            run = par.add_run(block[2:].replace("\n> ", " "))
            run.italic = True
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        elif block.startswith("- "):
            for item in block.split("\n"):
                add_runs(doc.add_paragraph(style="List Bullet"), item[2:])
        else:
            add_runs(doc.add_paragraph(), block.replace("\n", " "))

    doc.save(docx_path)
    print(f"saved {docx_path}")


if __name__ == "__main__":
    convert(sys.argv[1], sys.argv[2])
